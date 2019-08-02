from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def execute(*args):
	nombre=args[0]
	ap_pat=args[1]
	ap_mat=args[2]
	direccion=args[3]
	edad=args[4]
	tel=args[5]
	sexo=args[6]
	ocupacion=args[7]
	suceso=args[8]
	where=args[9]
	municipio=args[10]
	calle=args[11]
	hora=args[12]
	robados=args[13]
	document = Document()
	document.add_heading(' REPORTE DE DENUNCIA CDMX',0)
	p80 = document.add_paragraph("DENUNCIA PUBLICA").bold = True
	p0 = document.add_paragraph("DATOS DEL DENUNCIANTE")
	p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
	p = document.add_paragraph("Nombre del denunciante: ")
	p.add_run(nombre)
	p.add_run(ap_pat)
	p.add_run(ap_mat)
	p2 = document.add_paragraph("Direccion del denunciante: ")
	p2.add_run(direccion)
	p3 = document.add_paragraph("Edad del denunciante: ")
	p3.add_run(edad)
	p4 = document.add_paragraph("Telefono del denunciante: ")
	p4.add_run(tel)
	p5 = document.add_paragraph("Sexo: ")
	p5.add_run(sexo)
	p6 = document.add_paragraph("Ocupacion del denunciante: ")
	p6.add_run(ocupacion)
	p7 = document.add_paragraph("HISTORIA DE LOS HECHOS")
	p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
	p8 = document.add_paragraph("suceso: ")
	p8.add_run(suceso)
	p9 = document.add_paragraph("donde: ")
	p9.add_run(where)
	p10 = document.add_paragraph("Municipio: ")
	p10.add_run(municipio)
	p11 = document.add_paragraph("calle: ")
	p11.add_run(calle)
	p12 = document.add_paragraph("Hora: ")
	p12.add_run(hora)
	p13 = document.add_paragraph("OBJETOS PERDIDOS")
	p13.alignment = WD_ALIGN_PARAGRAPH.CENTER
	p14 = document.add_paragraph("objetos: ") 
	p14.add_run(robados)
	p15 = document.add_paragraph("DENUCIA HECHA A TRAVES DE DENUNCIABOT")
	p16 = document.add_paragraph("____________________________________________")
	document.add_page_break()
	document.save('REPORTE.docx')

