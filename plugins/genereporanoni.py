from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def execute(*args):
	suceso=args[0]
	where=args[1]
	municipio=args[2]
	calle=args[3]
	hora=args[4]
	robados=args[5]
	document = Document()
	document.add_heading('REPORTE DE DENUNCIA CDMX',0)
	p0 = document.add_paragraph("DENUNCIA ANONIMA").bold = True
	p1 = document.add_paragraph("HISTORIA DE LOS HECHOS")
	p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
	p2 = document.add_paragraph("suceso: ")
	p2.add_run(suceso)
	p3 = document.add_paragraph("donde: ")
	p3.add_run(where)
	p4 = document.add_paragraph("Municipio: ")
	p4.add_run(municipio)
	p5 = document.add_paragraph("calle: ")
	p5.add_run(calle)
	p6 = document.add_paragraph("Hora: ")
	p6.add_run(hora)
	p7 = document.add_paragraph("OBJETOS PERDIDOS")
	p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
	p8 = document.add_paragraph("objetos: ") 
	p8.add_run(robados)
	p9 = document.add_paragraph("DENUCIA HECHA A TRAVES DE DENUNCIABOT")
	p9 = document.add_paragraph("____________________________________________")
	document.add_page_break()
	document.save('REPORTE.docx')

